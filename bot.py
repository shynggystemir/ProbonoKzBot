import logging
from telegram import Update
from telegram.ext import (
    ApplicationBuilder,
    CommandHandler,
    ContextTypes,
    MessageHandler,
    filters,
    ConversationHandler,
)
from docx import Document

# Этапы диалога
FIO_IM, FIO_DAT, IIN, ADDRESS, PHONE, FIO_IM_DEBTOR, FIO_ROD_DEBTOR, IIN_DEBTOR, ADDRESS_DEBTOR, CHILD_NAME, CHILD_BIRTH, BIRTH_CONFIRMED, MARRIAGE_STATUS, DISTRICT = range(14)
user_data = {}

logging.basicConfig(format='%(asctime)s - %(name)s - %(levelname)s - %(message)s', level=logging.INFO)

async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text("Привет! Введи своё ФИО в именительном падеже:")
    return FIO_IM

async def get_input(update: Update, context: ContextTypes.DEFAULT_TYPE, key, prompt, next_state):
    if not update.message or not update.message.text:
        await update.message.reply_text("Пожалуйста, введи текст.")
        return next_state
    user_data[key] = update.message.text.strip()
    await update.message.reply_text(prompt)
    return next_state

async def get_fio_im(update, context): return await get_input(update, context, 'fio_im', "Теперь введи своё ФИО в дательном падеже:", FIO_DAT)
async def get_fio_dat(update, context): return await get_input(update, context, 'fio_dat', "Введи свой ИИН:", IIN)
async def get_iin(update, context): return await get_input(update, context, 'iin', "Укажи свой адрес (с индексом):", ADDRESS)
async def get_address(update, context): return await get_input(update, context, 'address', "Введи свой номер телефона:", PHONE)
async def get_phone(update, context): return await get_input(update, context, 'phone', "ФИО должника в именительном падеже:", FIO_IM_DEBTOR)
async def get_fio_im_debtor(update, context): return await get_input(update, context, 'fio_im_debtor', "ФИО должника в родительном падеже:", FIO_ROD_DEBTOR)
async def get_fio_rod_debtor(update, context): return await get_input(update, context, 'fio_rod_debtor', "ИИН должника (если известен):", IIN_DEBTOR)
async def get_iin_debtor(update, context): return await get_input(update, context, 'iin_debtor', "Адрес должника:", ADDRESS_DEBTOR)
async def get_address_debtor(update, context): return await get_input(update, context, 'address_debtor', "ФИО ребёнка:", CHILD_NAME)
async def get_child_name(update, context): return await get_input(update, context, 'child_name', "Дата рождения ребёнка:", CHILD_BIRTH)
async def get_child_birth(update, context): return await get_input(update, context, 'child_birth', "Указан ли должник в свидетельстве о рождении как отец? (да/нет)", BIRTH_CONFIRMED)

async def get_birth_confirmed(update, context):
    text = update.message.text.strip().lower()
    if text != "да":
        await update.message.reply_text("Если должник не указан в свидетельстве — подаётся иск, а не судебный приказ.")
        return ConversationHandler.END
    return await get_input(update, context, 'birth_confirmed', "Состоите ли вы в браке с должником? (состоим/развелись/не состояли)", MARRIAGE_STATUS)

async def get_marriage_status(update, context): return await get_input(update, context, 'marriage_status', "Укажите район проживания (Есильский, Сарыаркинский, Алматы, Байконур):", DISTRICT)

async def get_district(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_data['district'] = update.message.text.strip()

    # Генерация документа
    doc = Document()
    doc.add_heading('Заявление о выдаче судебного приказа', 0)
    doc.add_paragraph(f"От: {user_data['fio_dat']} (ИИН: {user_data['iin']})")
    doc.add_paragraph(f"Адрес: {user_data['address']}")
    doc.add_paragraph(f"Телефон: {user_data['phone']}")
    doc.add_paragraph("")
    doc.add_paragraph(f"Должник: {user_data['fio_im_debtor']} (ИИН: {user_data['iin_debtor']})")
    doc.add_paragraph(f"Адрес должника: {user_data['address_debtor']}")
    doc.add_paragraph("")
    doc.add_paragraph(f"Ребёнок: {user_data['child_name']}, {user_data['child_birth']}")
    doc.add_paragraph("")
    doc.add_paragraph("ЗАЯВЛЕНИЕ:")
    doc.add_paragraph("Прошу выдать судебный приказ о взыскании алиментов на содержание несовершеннолетнего ребёнка "
                      f"{user_data['child_name']} с должника {user_data['fio_rod_debtor']} в размере 1/4 от доходов ежемесячно, "
                      "начиная с даты подачи заявления и до достижения ребёнком совершеннолетия.")
    doc.save("/mnt/data/zayavlenie_ali_full.docx")

    await update.message.reply_text("Черновик заявления готов! Скачай его по ссылке: zayavlenie_ali_full.docx")
    return ConversationHandler.END

app = ApplicationBuilder().token("8160024624:AAFAI0alIhc6XdSc6WSB3LgKztxiDh2rYcE").build()

conv_handler = ConversationHandler(
    entry_points=[CommandHandler('start', start)],
    states={
        FIO_IM: [MessageHandler(filters.TEXT & ~filters.COMMAND, get_fio_im)],
        FIO_DAT: [MessageHandler(filters.TEXT & ~filters.COMMAND, get_fio_dat)],
        IIN: [MessageHandler(filters.TEXT & ~filters.COMMAND, get_iin)],
        ADDRESS: [MessageHandler(filters.TEXT & ~filters.COMMAND, get_address)],
        PHONE: [MessageHandler(filters.TEXT & ~filters.COMMAND, get_phone)],
        FIO_IM_DEBTOR: [MessageHandler(filters.TEXT & ~filters.COMMAND, get_fio_im_debtor)],
        FIO_ROD_DEBTOR: [MessageHandler(filters.TEXT & ~filters.COMMAND, get_fio_rod_debtor)],
        IIN_DEBTOR: [MessageHandler(filters.TEXT & ~filters.COMMAND, get_iin_debtor)],
        ADDRESS_DEBTOR: [MessageHandler(filters.TEXT & ~filters.COMMAND, get_address_debtor)],
        CHILD_NAME: [MessageHandler(filters.TEXT & ~filters.COMMAND, get_child_name)],
        CHILD_BIRTH: [MessageHandler(filters.TEXT & ~filters.COMMAND, get_child_birth)],
        BIRTH_CONFIRMED: [MessageHandler(filters.TEXT & ~filters.COMMAND, get_birth_confirmed)],
        MARRIAGE_STATUS: [MessageHandler(filters.TEXT & ~filters.COMMAND, get_marriage_status)],
        DISTRICT: [MessageHandler(filters.TEXT & ~filters.COMMAND, get_district)],
    },
    fallbacks=[]
)

app.add_handler(conv_handler)
app.run_polling()
